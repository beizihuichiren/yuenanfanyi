"""字幕审核校对模块。

提供两种校对方式：
  1. AI 二次校对（默认）：批量传入中文+越南语初译，AI 返回校对后的越南语
  2. 人工审核 DOCX（保留备用）：生成中越双语 DOCX 供人工填写意见
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Optional

from .pipeline_algorithms import SrtEntry, parse_srt, _time_to_seconds, SRT_TIME

logger = logging.getLogger(__name__)


def generate_review_docx(cn_entries: list[SrtEntry], vi_entries: list[SrtEntry],
                         output_path: Path) -> Path:
    """
    按序号对齐两份 SRT，生成 Word 表格。
    列：序号 | 时间轴 | 中文字幕 | 越南语初译 | 审核意见 | 修改建议
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError as exc:
        raise RuntimeError("python-docx 未安装，无法生成审核 DOCX") from exc

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"].font.name = "Arial"

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["序号", "时间轴", "中文字幕", "越南语初译", "审核意见", "修改建议"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = 1

    cn_map = {e.index: e for e in cn_entries}
    vi_map = {e.index: e for e in vi_entries}

    for idx in sorted(set(cn_map.keys()) | set(vi_map.keys())):
        cn = cn_map.get(idx)
        vi = vi_map.get(idx)
        row = table.add_row()
        row.cells[0].text = str(idx)
        if cn:
            row.cells[1].text = f"{_seconds_to_time(cn.start)} --> {_seconds_to_time(cn.end)}"
            row.cells[2].text = cn.text
        if vi:
            row.cells[3].text = vi.text

    widths = [Cm(1.2), Cm(3.5), Cm(4), Cm(4), Cm(3), Cm(3)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    doc.save(str(output_path))
    logger.info("审核文档已生成: %s", output_path)
    return output_path


def parse_review_docx(docx_path: Path) -> list[dict]:
    """
    从审核 DOCX 中提取有修改意见的行。
    返回 [{index, cn_text, vi_original, opinion, suggestion}, ...]
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx 未安装，无法解析审核 DOCX") from exc

    doc = Document(str(docx_path))
    items = []
    for table in doc.tables:
        for row in table.rows[1:]:
            cells = row.cells
            if len(cells) < 6:
                continue
            opinion = cells[4].text.strip()
            suggestion = cells[5].text.strip()
            if not opinion and not suggestion:
                continue
            idx_str = cells[0].text.strip()
            items.append({
                "index": int(idx_str) if idx_str.isdigit() else 0,
                "cn_text": cells[2].text.strip(),
                "vi_original": cells[3].text.strip(),
                "opinion": opinion,
                "suggestion": suggestion,
            })
    return items


def apply_review(vi_entries: list[SrtEntry], review_items: list[dict],
                 api_key: str, base_url: str, model: str) -> list[SrtEntry]:
    """
    根据审核意见，逐条/统一应用到越南语 SRT。
    通过 API 完成；无审核意见或 API 不可用时直接返回原 SRT。
    """
    if not review_items:
        logger.info("无审核意见，跳过 apply_review")
        return vi_entries

    if not api_key:
        logger.warning("未配置 API key，apply_review 直接返回初译")
        return vi_entries

    import requests

    srt_text = _entries_to_srt(vi_entries)
    review_text = "\n".join(
        f"序号{r['index']}: 原文={r['vi_original']}, 意见={r['opinion']}, 建议={r.get('suggestion', '')}"
        for r in review_items
    )

    prompt = f"""根据以下审核意见，对越南语 SRT 字幕进行全面修改。

审核意见：
{review_text}

修改要求：
1. 逐条落实上述审核意见
2. 对于同类问题（如人名/称谓不统一），检查整份字幕并统一修改
3. 完整保留序号、时间轴、条目数量，只修改文字

SRT 原文：
{srt_text}

请输出修改后的完整 SRT 文件。"""

    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": "你是专业越南语字幕校对员。"},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.2,
    }
    try:
        resp = requests.post(
            f"{base_url}/chat/completions",
            headers={"Authorization": f"Bearer {api_key}",
                     "Content-Type": "application/json"},
            json=payload, timeout=180,
        )
        resp.raise_for_status()
        content = resp.json()["choices"][0]["message"]["content"]
        updated = parse_srt_from_text(content)
        if len(updated) != len(vi_entries):
            logger.warning("apply_review 条目数不一致 (%d vs %d)，回退到初译",
                           len(updated), len(vi_entries))
            return vi_entries
        return updated
    except Exception as exc:
        logger.warning("apply_review 调用失败，回退到初译: %s", exc)
        return vi_entries


def _entries_to_srt(entries: list[SrtEntry]) -> str:
    blocks = []
    for e in entries:
        blocks.append(
            f"{e.index}\n"
            f"{_seconds_to_time(e.start)} --> {_seconds_to_time(e.end)}\n"
            f"{e.text}"
        )
    return "\n\n".join(blocks)


def _seconds_to_time(sec: float) -> str:
    h = int(sec // 3600)
    m = int((sec % 3600) // 60)
    s = int(sec % 60)
    ms = int((sec - int(sec)) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def parse_srt_from_text(text: str) -> list[SrtEntry]:
    """从纯文本解析 SRT，兼容 API 返回格式"""
    entries = []
    blocks = text.strip().split("\n\n")
    for block in blocks:
        lines = block.strip().split("\n")
        if len(lines) < 3:
            continue
        try:
            idx = int(lines[0].strip())
        except ValueError:
            continue
        times = re.findall(SRT_TIME, lines[1])
        if len(times) != 2:
            continue
        start = _time_to_seconds(times[0])
        end = _time_to_seconds(times[1])
        txt = "\n".join(lines[2:])
        entries.append(SrtEntry(index=idx, start=start, end=end, text=txt))
    return entries


# ==================== AI 二次校对（默认校对方式） ====================

AI_REVIEW_SYSTEM_PROMPT = """你是专业的越南语字幕校对员。你会收到中文原文和越南语初译，
请对照中文原文对越南语译文进行二次校对，输出校对后的越南语字幕。

校对要点：
1. 准确性：修正错译、漏译、多译
2. 流畅性：使越南语表达自然地道，符合口语习惯
3. 一致性：统一人名、称谓、语气词的译法（整份字幕保持一致）
4. 时长控制：译文长度尽量与原文相近，避免过长导致字幕显示不全
5. 保留格式：完整保留序号、时间轴、条目数量，只修改译文文字

输出格式：完整的 SRT 文件内容，不附加任何解释说明。"""


def ai_review_batch(cn_entries: list[SrtEntry],
                    vi_entries: list[SrtEntry],
                    api_key: str,
                    base_url: str,
                    model: str,
                    batch_size: int = 40) -> list[SrtEntry]:
    """
    AI 二次校对：对照中文原文，批量校对越南语初译。

    分批调用 API（每批 batch_size 条），避免单次请求过长。
    校对失败时该批回退到初译，不影响整体流程。

    Args:
        cn_entries:    中文字幕条目
        vi_entries:    越南语初译条目（序号/时间轴与 cn 对齐）
        api_key:       OpenAI 兼容 API key
        base_url:      API base url
        model:         模型名
        batch_size:    每批条目数

    Returns:
        校对后的越南语条目列表（序号/时间轴保持不变）
    """
    import requests

    if not api_key:
        logger.warning("未配置 API key，AI 二次校对跳过，使用初译版")
        return vi_entries

    if len(cn_entries) != len(vi_entries):
        logger.warning("中越条目数不一致 (%d vs %d)，AI 校对跳过",
                       len(cn_entries), len(vi_entries))
        return vi_entries

    # 按序号对齐
    cn_map = {e.index: e for e in cn_entries}
    vi_map = {e.index: e for e in vi_entries}
    all_indices = sorted(cn_map.keys())

    reviewed: list[SrtEntry] = []
    total_batches = (len(all_indices) + batch_size - 1) // batch_size

    for batch_idx in range(total_batches):
        start_i = batch_idx * batch_size
        end_i = min(start_i + batch_size, len(all_indices))
        batch_indices = all_indices[start_i:end_i]

        # 构造批次输入：序号 | 中文 | 越南语初译
        batch_lines = []
        for idx in batch_indices:
            cn = cn_map[idx]
            vi = vi_map[idx]
            batch_lines.append(
                f"[{idx}]\n中文: {cn.text}\n越译: {vi.text}"
            )
        batch_input = "\n\n".join(batch_lines)

        prompt = f"""请校对以下 {len(batch_indices)} 条字幕的越南语译文。

输入（每条包含序号、中文原文、越南语初译）：
{batch_input}

请输出校对后的完整 SRT（保留原序号和时间轴），格式如下：
序号
00:00:00,000 --> 00:00:00,000
校对后的越南语译文

时间轴参考：
""" + "\n".join(
            f"{vi_map[idx].index}\n"
            f"{_seconds_to_time(vi_map[idx].start)} --> {_seconds_to_time(vi_map[idx].end)}\n"
            f"{vi_map[idx].text}"
            for idx in batch_indices
        )

        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": AI_REVIEW_SYSTEM_PROMPT},
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.2,
        }

        try:
            resp = requests.post(
                f"{base_url}/chat/completions",
                headers={"Authorization": f"Bearer {api_key}",
                         "Content-Type": "application/json"},
                json=payload, timeout=120,
            )
            resp.raise_for_status()
            content = resp.json()["choices"][0]["message"]["content"]
            batch_reviewed = _parse_reviewed_batch(content, batch_indices, vi_map)
            reviewed.extend(batch_reviewed)
            logger.info("AI 校对批次 %d/%d 完成 (%d 条)",
                        batch_idx + 1, total_batches, len(batch_reviewed))
        except Exception as exc:
            logger.warning("AI 校对批次 %d/%d 失败，回退初译: %s",
                           batch_idx + 1, total_batches, exc)
            reviewed.extend(vi_map[idx] for idx in batch_indices)

    # 按序号排序输出
    reviewed.sort(key=lambda e: e.index)
    return reviewed


def _parse_reviewed_batch(text: str, expected_indices: list[int],
                          vi_map: dict[int, SrtEntry]) -> list[SrtEntry]:
    """
    从 AI 返回的文本中解析校对后的条目。
    解析失败或条目数不匹配时回退到初译。
    """
    parsed = parse_srt_from_text(text)
    parsed_map = {e.index: e for e in parsed}

    result = []
    for idx in expected_indices:
        if idx in parsed_map:
            # 保留原时间轴（AI 可能改坏），只用校对后的文字
            orig = vi_map[idx]
            result.append(SrtEntry(
                index=orig.index,
                start=orig.start,
                end=orig.end,
                text=parsed_map[idx].text,
            ))
        else:
            result.append(vi_map[idx])
    return result
